from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_document():
    # Создаём документ
    document = Document()

    # Заголовок
    title = document.add_paragraph("В микроконтроллерах ATmega, используемых на платформах Arduino, существует три вида памяти:")
    title_format = title.paragraph_format
    title_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(12)

    # Описание памяти
    document.add_paragraph("Флеш-память: используется для хранения скетчей.", style="List Bullet")

    bullet2 = document.add_paragraph("ОЗУ (", style="List Bullet")
    bullet2.add_run("SRAM").bold = True
    bullet2.add_run(" — ")
    bullet2.add_run("static random access memory").italic = True
    bullet2.add_run(", статическая оперативная память с произвольным доступом): используется для хранения переменных.")

    document.add_paragraph("EEPROM (энергонезависимая память): используется для хранения постоянной информации.", style="List Bullet")

    document.add_paragraph("Флеш-память и EEPROM являются энергонезависимыми видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью.")

    # Таблица
    table = document.add_table(rows=4, cols=5)
    table.style = 'Table Grid'

    contents = (
        ("", "ATmega168", "ATmega328", "ATmega1280", "ATmega2560"),
        ("Flash (1 кБ flash-памяти занят загрузчиком)", "16 Кбайт", "32 Кбайт", "128 Кбайт", "256 Кбайт"),
        ("SRAM", "1 Кбайт", "2 Кбайт", "8 Кбайт", "8 Кбайт"),
        ("EEPROM", "512 байт", "1024 байта", "4 Кбайт", "4 Кбайт")
    )

    headers = table.rows[0].cells
    for i in range(5):
        header = headers[i].add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.add_run(contents[0][i]).bold = True

    entries = table.columns[0].cells
    for i in range(1, 4):
        entry = entries[i].add_paragraph()
        entry.alignment = WD_ALIGN_PARAGRAPH.CENTER
        entry.add_run(contents[i][0]).bold = True

    for row in range(1, 4):
        for col in range(1, 5):
            data = table.cell(row, col).add_paragraph()
            data.alignment = WD_ALIGN_PARAGRAPH.CENTER
            data.add_run(contents[row][col])

    document.add_paragraph()

    # Последний абзац
    last = document.add_paragraph()
    last.add_run("Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25°C. Эти данные не распространяются на операции чтения данных из EEPROM — чтение данных не лимитированно. Исходя из этого, нужно проектировать свои скетчи максимально щадящими по отношению к EEPROM.").italic = True

    # Сохранение документа
    document.save("Task1.docx")


if __name__ == "__main__":
    create_document()
