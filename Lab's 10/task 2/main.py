from docx import Document
from docx.shared import Cm

document = Document("Task1.docx")

document.add_picture("unnamed.jpg", width=Cm(5), height=Cm(5))
document.add_paragraph("Ядер кола")

document.save("Task1.docx")